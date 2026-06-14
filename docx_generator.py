from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Brand colors
NAVY = RGBColor(0x0B, 0x1F, 0x3A)
TEAL = RGBColor(0x00, 0xA9, 0x9D)


def generate_docx(data: dict) -> bytes:
    """
    Takes the report data dictionary from cms_client.py
    and builds a Word document using python-docx.
    Returns the docx as bytes so FastAPI can stream it.
    """

    doc = Document()

    # ── Page Margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # ── Header — Platform Branding ────────────────────────────
    # INFINITE branding — always hardcoded, never dynamic
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(data["platform"])
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.color.rgb = TEAL

    # ── Report Title ──────────────────────────────────────────
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(data["report_title"])
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.color.rgb = NAVY

    # ── State Code ────────────────────────────────────────────
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(data["state"])
    run3.bold = True
    run3.font.size = Pt(13)
    run3.font.color.rgb = NAVY

    doc.add_paragraph()

    # ── Data Table ────────────────────────────────────────────
    # All 26 fields in a two column table
    fields = [
        ("Name of Facility",                            data["facility_name"]),
        ("Location",                                    data["location"]),
        ("EMR",                                         data["emr"]),
        ("Census Capacity",                             str(data["census_capacity"])),
        ("Current Census",                              data["current_census"]),
        ("Type of Patient",                             data["patient_type"]),
        ("Previous Coverage from Medelite",             data["previous_coverage"]),
        ("Previous Provider Performance from Medelite", data["previous_performance"]),
        ("Medical Coverage",                            data["medical_coverage"]),
        ("Overall Star Rating",                         str(data["overall_rating"])),
        ("Health Inspection",                           str(data["health_inspection_rating"])),
        ("Staffing",                                    str(data["staffing_rating"])),
        ("Quality of Resident Care",                    str(data["quality_measure_rating"])),
        ("Short Term Hospitalization",                  data["str_hospitalization"]),
        ("STR National Avg. for Hospitalization",       data["str_hosp_national_avg"]),
        ("STR State Avg. for Hospitalization",          data["str_hosp_state_avg"]),
        ("STR ED Visit",                                data["str_ed_visit"]),
        ("STR ED Visits National Avg.",                 data["str_ed_national_avg"]),
        ("STR ED Visits State Avg.",                    data["str_ed_state_avg"]),
        ("LT Hospitalization",                          data["lt_hospitalization"]),
        ("LT National Avg. for Hospitalization",        data["lt_hosp_national_avg"]),
        ("LT State Avg. for Hospitalization",           data["lt_hosp_state_avg"]),
        ("ED Visit",                                    data["lt_ed_visit"]),
        ("LT ED Visits National Avg.",                  data["lt_ed_national_avg"]),
        ("LT ED Visits State Avg.",                     data["lt_ed_state_avg"]),
        ("Medicare Source",                             data["medicare_url"]),
    ]

    # Create the table with a header row
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"

    # Style the header row
    for cell in hdr:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = NAVY
        run.font.size = Pt(10)

    # Add all data rows
    for label, value in fields:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

        # Make the label bold
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].paragraphs[0].runs[0].font.size = Pt(9)

    # ── Build and Return ──────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()